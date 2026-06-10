# -*- coding: utf-8 -*-
"""Word report: 3 designs per question, each with strengths & weaknesses."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
ACCENT2 = RGBColor(0x2E, 0x75, 0xB6)


def h1(text):
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(15); r.font.color.rgb = ACCENT


def design_title(text):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = ACCENT2


def add_image(path, width=4.8):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"[missing image: {path}]")


def bullets(title, items):
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def design(title, desc, image, width, strengths, weaknesses):
    design_title(title)
    doc.add_paragraph(desc)
    add_image(image, width)
    bullets("Strengths", strengths)
    bullets("Weaknesses", weaknesses)


# ---------------- Title ----------------
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Baby Names in France (1900–2020)")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = ACCENT
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = s.add_run("Design Alternatives — Strengths & Weaknesses")
rs.italic = True; rs.font.size = Pt(13)
doc.add_paragraph(
    "For each set of questions we present three design alternatives, each with its strengths and "
    "weaknesses. The guiding principle is overview + detail on demand: one interactive "
    "visualization should answer all sub-questions. Popularity is normalized as a share of births "
    "(%) rather than raw counts, so years and departments are comparable."
)

# ============================================================
# VIZ 1
# ============================================================
h1("Visualization 1 — Evolution over time")
doc.add_paragraph(
    "Questions: which names stay consistently (un)popular, which are suddenly or briefly popular, "
    "and are there temporal trends?"
)

design(
    "Design 1.1 — Diagonal heatmap (names × years, sorted by peak year)",
    "Each row is a name, each column a year, color encodes the name's yearly share. Rows are "
    "ordered by their peak year, producing a diagonal that makes every name's era visible at once.",
    "sketches/WhatsApp Image 2026-06-06 at 23.53.00.jpeg", 3.4,
    ["The diagonal reveals trends instantly: consistency (long horizontal bands), brief fads "
     "(short spots) and accelerating turnover toward recent decades.",
     "Shows many names simultaneously without clutter.",
     "Encodes magnitude (share) directly through color."],
    ["Static; reading one specific name precisely is hard.",
     "The Top-N of names to display must be chosen.",
     "Subtle differences can be lost in the color scale."],
)

design(
    "Design 1.2 — Longevity quadrant + linked time series",
    "A scatter plot positions each name by peak year (X) and longevity — years spent above a "
    "popularity threshold (Y); bubble size = total births. Clicking a bubble shows its full time "
    "series below. Names sort into classics (top), fads (bottom), older/recent eras (left/right).",
    "sketches/viz1-nuage_de_bulles_v2.png", 6.0,
    ["Directly answers 'consistent vs sudden/brief' — it is the chart's vertical axis.",
     "Overview + detail in one view via the linked time series on click.",
     "Bubble size keeps absolute magnitude visible, not just rank."],
    ["Bubble overplotting in dense areas needs zoom/jitter/opacity.",
     "'Longevity' relies on an arbitrary threshold that must be stated.",
     "Locating a given name still requires interaction (search/hover)."],
)

design(
    "Design 1.3 — Bump chart / 'metro' lines (Top-10 rank over time)",
    "Lines track the yearly rank of the Top-10 names; crossings show overtakes and rank reversals "
    "across the century.",
    "sketches/WhatsApp Image 2026-06-06 at 23.52.57.jpeg", 5.6,
    ["Excellent for crossings, comebacks and rank reversals.",
     "Very legible for a small set of leaders.",
     "Intuitive narrative of 'who was on top when'."],
    ["Limited to ~10 names — misses obscure or briefly popular names.",
     "Rank hides magnitude (the #1 name may be far ahead of #2).",
     "Becomes a tangle if too many names are shown."],
)

# ============================================================
# VIZ 2
# ============================================================
h1("Visualization 2 — Regional effect")
doc.add_paragraph(
    "Questions: are some names more popular in certain regions, and are popular names generally "
    "popular across the whole country?"
)

design(
    "Design 2.1 — Interactive choropleth of the location quotient (LQ)",
    "A department map colored by a selected name's location quotient = local share ÷ national "
    "share. LQ > 1 (red) = over-represented locally; ~1 = national average. A name selector and a "
    "year slider let the user explore any name and watch its geographic diffusion over time.",
    "sketches/viz2_A_choropleth.png", 4.4,
    ["Geographic encoding is the natural answer to a geographic question.",
     "The LQ isolates the regional effect and corrects for department size.",
     "The year slider exposes diffusion (names often spread from a region outward)."],
    ["Shows one name at a time; cross-name comparison needs another view.",
     "Insee codes Corsica as '20' while the map uses 2A/2B (reconciliation needed).",
     "Overseas departments are off-scale and excluded (metropolitan France only)."],
)

design(
    "Design 2.2 — 'Signature name' per department map",
    "Each department is labeled with its most over-represented name (highest LQ), exposing "
    "regional cultural identities — Breton, Basque, Catalan, Corsican signatures.",
    "sketches/viz2_C_signature.png", 4.6,
    ["Reveals every region's identity in a single static picture.",
     "Strong, memorable storytelling of regional culture.",
     "Directly answers 'are some names more popular in some regions?'."],
    ["96 labels make the map crowded.",
     "Sensitive to outliers and to the minimum-births threshold.",
     "Shows only the single top name, not the broader distribution."],
)

design(
    "Design 2.3 — Sunburst (region → department hierarchy)",
    "Concentric rings break down births by region (inner ring) and department (outer ring), "
    "showing how each region is composed and the relative weight of its departments.",
    "sketches/viz2-3.png", 4.2,
    ["Clear hierarchical composition: region totals and their department breakdown together.",
     "Compact part-to-whole view; interactive drill-down is natural.",
     "Good complement to the map for 'how big is each region'."],
    ["No geographic dimension — cannot answer 'where' on the territory.",
     "Hard to compare angular slice sizes precisely.",
     "Does not express 'local vs national' popularity by itself (secondary view)."],
)

# ============================================================
# VIZ 3
# ============================================================
h1("Visualization 3 — Gender effect")
doc.add_paragraph(
    "Questions: are there gender effects, and does the popularity of names given to both sexes "
    "evolve consistently across sexes?"
)

design(
    "Design 3.1 — Female-share heatmap per decade",
    "Rows are unisex names, columns are decades, color is the female share on a diverging scale "
    "(blue = mostly male, orange = mostly female, grey = balanced). Gender flips appear as a row "
    "changing color over time.",
    "sketches/WhatsApp Image 2026-06-07 at 19.35.20 (3).jpeg", 4.6,
    ["Surfaces gender flips (e.g. Camille, Alix) across many names at once.",
     "Diverging color makes 'consistent vs shifting' legible at a glance.",
     "Compact overview of the whole set of unisex names."],
    ["Decade aggregation smooths short-lived swings.",
     "Limited to names given to both sexes.",
     "Color alone cannot show the volume behind each share."],
)

design(
    "Design 3.2 — Female-share line chart (interactive, click to isolate)",
    "One line per unisex name shows its female share year by year; clicking the legend isolates a "
    "name (Shift+click for several) for precise temporal reading.",
    "sketches/WhatsApp Image 2026-06-07 at 19.35.20.jpeg", 5.2,
    ["Year-level precision on each name's gender trajectory.",
     "Interaction (isolate/compare) tames the complexity on demand.",
     "Crossing the 50% line marks the exact moment a name flips."],
    ["Illegible spaghetti without default filtering to a few names.",
     "Hard to get an overview of all names at once.",
     "Overlapping lines obscure simultaneous trends."],
)

design(
    "Design 3.3 — Split male/female violins per name",
    "For each unisex name, a split violin shows the distribution of births over years for males "
    "vs females, revealing each name's active periods by sex.",
    "sketches/WhatsApp Image 2026-06-07 at 19.35.20 (1).jpeg", 6.0,
    ["Shows each name's active period and its male/female balance together.",
     "Distribution shape conveys more than a single share value.",
     "Visually distinctive and compact across many names."],
    ["Answers 'when' more than 'does the gender flip'.",
     "Violin shapes are harder for a general audience to read.",
     "Comparing many names side by side becomes dense."],
)

# ---- closing ----
h1("Cross-cutting consistency")
doc.add_paragraph(
    "Across all views we keep a consistent sex color encoding, reuse a shared 'heatmap' visual "
    "language (Viz 1 and Viz 3), and clean the data identically (rare names, dpt=XX and "
    "annais=XXXX excluded)."
)

out = "report/Baby_Names_Designs_Report.docx"
doc.save(out)
print("Saved", out)
