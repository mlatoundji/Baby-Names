# -*- coding: utf-8 -*-
"""Final submission post (English, Word): the 3 implemented visualizations + rationale."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
ACCENT2 = RGBColor(0x2E, 0x75, 0xB6)


def h1(text):
    p = doc.add_paragraph(); p.space_before = Pt(14)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = ACCENT


def sub(text):
    p = doc.add_paragraph(); p.space_before = Pt(6)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = ACCENT2


def img(path, width=6.0):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run("Static snapshot — the live HTML adds the slider, hover and linked views.")
        r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    else:
        doc.add_paragraph(f"[missing image: {path}]")


def bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


# ---------------- Title ----------------
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Baby Names in France (1900–2020)")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = ACCENT
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = s.add_run("Final Visualizations — Presentation & Rationale")
rs.italic = True; rs.font.size = Pt(13)
doc.add_paragraph(
    "We implemented three interactive visualizations in Altair (Vega-Lite), one per set of "
    "questions. Each is a self-contained HTML file with a slider, tooltips and linked views. "
    "Popularity is measured as a share of births so that years, regions and sexes are comparable. "
    "Below, each visualization is presented with a snapshot and an explanation of why it is "
    "appropriate and effective for its target questions. (Screenshots here are static; the live "
    "files reveal the full interaction — we recommend a short screen-capture clip per "
    "visualization showing the slider being dragged.)"
)

# ============================================================
# VIZ 1
# ============================================================
h1("Visualization 1 — Evolution over time")
doc.add_paragraph(
    "Target questions: which names stay consistently (un)popular, which are suddenly or briefly "
    "popular, and are there temporal trends?"
)
sub("What it shows")
doc.add_paragraph(
    "An animated bubble cloud. Each bubble is a name in the selected year: the X-axis is the "
    "name's peak year, the Y-axis is its longevity (the number of years it stayed in the national "
    "Top 100), and bubble area is the number of births that year, on a fixed scale so sizes stay "
    "comparable from one year to the next. Names are pooled regardless of sex — one bubble per "
    "name — since the questions concern names, not the child's sex. The year selector is the "
    "X-axis itself: a handle slides along the time axis (or the ▶ button plays the years in "
    "sequence), while the giant year label and the moving, growing bubbles make the evolution "
    "tangible and a dashed ghost ring recalls each bubble's size the year before. Two filters "
    "sharpen the reading — a “Top n” slider (1–100) thins the cloud from the headline names down "
    "to the long tail, and a “Follow a name” box highlights one name, fading the rest and drawing "
    "its full 1900–2020 births curve above the chart."
)
img("sketches/viz1_bulles_preview.png", 6.2)
sub("Why it is appropriate and effective")
bullets([
    "Consistent vs brief popularity is read directly off the vertical axis: long-lived classics "
    "sit high (e.g. Marie), short-lived fads sit low — answering the core question by position.",
    "Following a name draws its births curve over the whole century, so the shape settles the "
    "question at a glance: a broad plateau is an enduring favourite, a single narrow spike a "
    "passing fad (e.g. Brigitte around 1960, Kevin around 1990).",
    "Temporal trends emerge from scrubbing: the cloud drifts rightward over the century and names "
    "are replaced faster in recent decades, making accelerating turnover visible.",
    "Sudden popularity is shown by a bubble that appears and inflates rapidly as the slider "
    "advances, while its low longevity keeps it near the bottom.",
    "The “Top n” filter controls clutter: a low n isolates the dominant names, a high n exposes "
    "the churn of the long tail without changing the layout.",
    "Hover tooltips give the exact name, births, rank within the year, peak year and longevity "
    "for detail on demand.",
])
sub("Key points to demonstrate (screenshots / video)")
bullets([
    "Snapshots at 1905, 1965 and 2015 to show how the set of names turns over.",
    "A long-lived name high on the axis (Marie/Jean) versus a brief fad low on the axis.",
    "“Follow a name” on a fad to reveal its sharp, isolated births curve — e.g. Brigitte (a 1960s "
    "spike) or Camille (a telling bimodal curve once its two eras are pooled).",
    "A short clip dragging the slider along the axis to show the cloud shifting and bubbles growing.",
])

# ============================================================
# VIZ 2
# ============================================================
h1("Visualization 2 — Regional effect")
doc.add_paragraph(
    "Target questions: are some names more popular in some regions, and are popular names "
    "generally popular across the whole country?"
)
doc.add_paragraph(
    "We use two complementary views: a per-name radial chart (the main answer) and a sunburst "
    "(demographic context).", style=None)

sub("Main view — radial chart of a name's local popularity")
doc.add_paragraph(
    "Each department is a fixed-width angular sector, arranged in pseudo-geographic order "
    "(west on the left, east on the right, north on top), so the layout is always the same and "
    "recognizable. The radial length of a sector encodes the name's local popularity as a "
    "frequency — births per thousand in that department — normalized by the national rate, i.e. "
    "how many times more (or less) common the name is locally than nationally. A fixed dashed "
    "reference circle marks the national level (×1): a sector reaching beyond the circle means the "
    "name is over-represented there, inside means under-represented. A name selector (sorted by "
    "popularity) and a decade slider drive the chart; region codes label the ring and are repeated "
    "with full names in the legend."
)
img("sketches/viz2_sunburst_2_demo_ewen.png", 6.0)
sub("Why it is appropriate and effective")
bullets([
    "Frequency-normalized length removes the population-size bias: Paris (75) is no longer "
    "automatically ‘biggest’; we compare like with like.",
    "The fixed ×1 reference circle turns the regional question into a one-glance read: bars beyond "
    "the circle = locally more popular, bars inside = less.",
    "It answers both sub-questions directly — a regional name (Breton EWEN) makes the western "
    "sectors shoot past the circle while everywhere else stays inside; a top national name (EMMA) "
    "produces near-uniform bars hugging the circle everywhere (popular across the whole country).",
    "The stable geographic ordering lets the eye learn the layout once and then compare any name; "
    "the decade slider shows how a name's geography appears, spreads or fades over time.",
])

sub("Complementary view — sunburst of births by region and department")
doc.add_paragraph(
    "The sunburst (inner ring = region, outer ring = departments, angle = number of births, decade "
    "slider) gives the demographic context the radial chart abstracts away: how many births each "
    "region actually contributes, and how that weight shifts across decades. It is the "
    "‘how big is each region’ overview; the radial chart is the ‘where is this name "
    "over-represented’ detail."
)
img("sketches/viz2_sunburst_1_preview.png", 4.4)
sub("Key points to demonstrate (screenshots / video)")
bullets([
    "Radial chart with a regional name (EWEN/JORDI/MAYLIS) — western/southern sectors past the circle.",
    "Radial chart with a national name (EMMA/LUCAS) — uniform bars on the reference circle.",
    "A short clip changing the name in the selector to show the pattern reshaping.",
])

# ============================================================
# VIZ 3
# ============================================================
h1("Visualization 3 — Gender effect")
doc.add_paragraph(
    "Target questions: are there gender effects, and does the popularity of names given to both "
    "sexes evolve consistently across sexes?"
)
sub("What it shows")
doc.add_paragraph(
    "Split violin plots — one per unisex name. Each violin shows the distribution of births over "
    "time (vertical axis = year): the left half (blue) is boys, the right half (orange) is girls, "
    "and the width at each year is the number of births that year for that sex. Names are ordered "
    "from the most male-leaning to the most female-leaning, and hovering a name highlights its "
    "violin."
)
img("sketches/viz3_violon_preview.png", 6.4)
sub("Why it is appropriate and effective")
bullets([
    "It focuses on unisex names — the only names for which a gender question is meaningful — "
    "so the view is on-topic by construction.",
    "A gender flip is read directly from the shape: when the blue and orange halves are offset "
    "vertically, the name was given to one sex in one era and the other sex later (e.g. Camille — "
    "a little blue early, a large orange bulge recently).",
    "Consistent evolution shows as two halves that rise and fall together over the same years; an "
    "inconsistent / shifting name shows halves peaking in different decades.",
    "Putting all names side by side, ordered by overall female share, turns the whole set of "
    "unisex names into one comparable small-multiples panel.",
    "Width = births per year, so each violin also conveys when the name was actually in use, not "
    "just its balance.",
])
sub("Key points to demonstrate (screenshots / video)")
bullets([
    "The full panel, showing names ranging from male-dominated (left) to female-dominated (right).",
    "A flipped name (Camille, Alix) whose blue and orange halves sit in different eras.",
    "A name used for both sexes in the same period (Dominique) as a balanced counter-example.",
])

# ---- closing ----
h1("Summary")
doc.add_paragraph(
    "Each visualization is matched to its question through a deliberate encoding: longevity on an "
    "axis for time, a frequency-normalized radial chart with a national reference circle (plus a "
    "sunburst for demographic context) for region, and split male/female violins over time for "
    "gender. All three follow the same overview + detail-on-demand pattern (slider / hover / "
    "linked view) and share a consistent, accessible color language."
)

out = "report/Final_Visualizations_Report.docx"
doc.save(out)
print("Saved", out)
